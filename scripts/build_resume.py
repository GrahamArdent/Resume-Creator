from __future__ import annotations
import json, sys
from pathlib import Path
from docx import Document
from utils import (
    extract_keywords, md_experience, fill_template, to_ascii,
    slugify, ProfileSchema, AnswersSchema, dedupe_list
)
from scripts.prompt_engine import rank_keywords_with_llm

BASE = Path(__file__).resolve().parents[1]

def _write_markdown_as_docx(md: str, profile: dict, out_path: Path):
    doc = Document()
    doc.add_paragraph(profile["name"])
    doc.add_paragraph(f"{profile['location']} | {profile['email']} | {profile['phone']} | {profile['linkedin']}")
    doc.add_paragraph("")
    for raw in (md or "").splitlines():
        line = raw.rstrip()
        if not line:
            doc.add_paragraph("")
            continue
        doc.add_paragraph(to_ascii(line))
    doc.save(out_path)

def _lint_docx(path: Path) -> list[str]:
    issues = []
    try:
        doc = Document(path)
        if getattr(doc, "tables", []):
            if len(doc.tables) > 0:
                issues.append(f"Found {len(doc.tables)} table(s) — remove tables for ATS.")
        try:
            sec = doc.sections[0]
            if any(p.text.strip() for p in sec.header.paragraphs):
                issues.append("Header contains text — clear it (ATS can skip headers).")
            if any(p.text.strip() for p in sec.footer.paragraphs):
                issues.append("Footer contains text — clear it (ATS can skip footers).")
        except Exception:
            pass
        full = "\n".join(p.text for p in doc.paragraphs)
        bad = [ch for ch in full if ord(ch) > 127]
        if bad:
            issues.append(f"Non-ASCII chars detected: {len(bad)} — keep ASCII only.")
    except Exception as e:
        issues.append(f"Lint error: {e}")
    return issues

def _load_profile() -> dict:
    data = json.loads((BASE / "profile" / "profile.json").read_text(encoding="utf-8"))
    ProfileSchema(**data)
    return data

def _load_answers() -> dict:
    p = BASE / "profile" / "answers.json"
    if p.exists():
        data = json.loads(p.read_text(encoding="utf-8"))
    else:
        data = {"global":{"extra_keywords":[],"summary_additions":[]},"roles":{},"asked":{}}
    try:
        AnswersSchema.model_validate(data)
    except Exception:
        pass
    return data

def _compose_markdown(profile: dict, answers: dict, jd_text: str, style: str):
    kw_candidates = extract_keywords(jd_text)
    ranked = rank_keywords_with_llm(jd_text, kw_candidates)

    g = answers.get("global", {})
    extras = g.get("summary_additions", [])
    summary_bits = [
        "Enterprise Account Executive for cybersecurity/SaaS across enterprise & public sector",
        "Full sales cycle: prospecting, discovery, negotiation, close (MEDDIC, BANT)",
        "CRM discipline (Salesforce/HubSpot); executive storytelling for identity-first, passwordless, endpoint-enforced access",
    ] + extras
    summary_bits = [to_ascii(s) for s in summary_bits if s]
    summary = "; ".join(summary_bits).rstrip(";") + "."

    def filter_join(items: list[str], cap: int = 8) -> str:
        keep = []
        for it in items:
            if any(it.lower() in rk.lower() for rk in ranked):
                keep.append(it)
        if not keep: keep = items[:cap]
        return ", ".join(dedupe_list(keep))[:400]

    domains = filter_join(profile.get("domains", []), 6)
    methods = filter_join(profile.get("methods", []), 10)
    platforms = filter_join(profile.get("platforms", []), 8)
    extra_kw = g.get("extra_keywords", [])
    security = ", ".join(dedupe_list([*profile.get("security_terms", []), *extra_kw]))[:500]
    collab = "Marketing, Presales, Leadership; Partner co-selling; Playbook mentoring"

    exp_md = md_experience(profile["experience"], answers)
    education = "\n".join(f"- {to_ascii(e)}" for e in profile.get("education", [])) or "-"
    awards = "\n".join(f"- {to_ascii(a)}" for a in profile.get("awards", [])) or "-"

    tpl_map = {"executive":"executive.md","balanced":"balanced.md","ats":"ats_strict.md","human":"human.md","cv":"cv.md"}
    tpl = (BASE / "templates" / tpl_map.get(style, "balanced.md")).read_text(encoding="utf-8")

    md = fill_template(tpl, {
        "NAME": profile["name"],
        "LOCATION": profile["location"],
        "LINKEDIN": profile["linkedin"],
        "EMAIL": profile["email"],
        "PHONE": profile["phone"],
        "SUMMARY": summary,
        "DOMAINS": domains,
        "METHODS": methods,
        "PLATFORMS": platforms,
        "SECURITY": security,
        "COLLAB": collab,
        "EXPERIENCE": exp_md,
        "EDUCATION": education,
        "AWARDS": awards
    })
    return md, ranked

def _write_all(md: str, ranked: list[str], profile: dict, out_root: Path, style_label: str):
    out_root.mkdir(parents=True, exist_ok=True)
    out_docx = out_root / f"Resume - {profile['name']} - {style_label} ({out_root.name}).docx"
    out_report = out_root / f"match_report_{style_label.lower()} ({out_root.name}).json"
    out_lint = out_root / f"ats_lint ({out_root.name}).txt"

    _write_markdown_as_docx(md, profile, out_docx)

    resume_text = "\n".join(p.text for p in Document(out_docx).paragraphs).lower()
    present = sorted([k for k in ranked if k.lower() in resume_text])
    missing = sorted([k for k in ranked if k.lower() not in resume_text])
    report = {
        "jd_tokens": ranked,
        "present": present,
        "missing": missing,
        "coverage_percent": round(100 * len(present) / max(1, len(ranked)), 1),
    }
    out_report.write_text(json.dumps(report, indent=2), encoding="utf-8")

    issues = _lint_docx(out_docx)
    out_lint.write_text("\n".join(issues) if issues else "No ATS lint issues detected.", encoding="utf-8")

    return {"docx": str(out_docx), "report": str(out_report), "lint": str(out_lint)}

def build_pair(primary_style: str = "balanced", company_slug: str = "generic", jd_text: str | None = None):
    profile = _load_profile()
    answers = _load_answers()
    if jd_text is None:
        jd_file = BASE / "data" / "job_posting.txt"
        jd_text = jd_file.read_text(encoding="utf-8") if jd_file.exists() else ""
    company_slug = slugify(company_slug)
    out_root = BASE / "outputs" / company_slug

    md_primary, ranked_primary = _compose_markdown(profile, answers, jd_text, style=primary_style)
    primary = _write_all(md_primary, ranked_primary, profile, out_root, style_label=primary_style.capitalize())

    md_cv, ranked_cv = _compose_markdown(profile, answers, jd_text, style="cv")
    cv = _write_all(md_cv, ranked_cv, profile, out_root, style_label="CV")

    return {"primary": primary, "cv": cv}

def main():
    primary_style = (sys.argv[1] if len(sys.argv)>1 else "balanced").lower()
    company_slug = sys.argv[2] if len(sys.argv)>2 else "generic"
    jd = None
    if len(sys.argv)>3:
        p = Path(sys.argv[3])
        if p.exists(): jd = p.read_text(encoding="utf-8")
    result = build_pair(primary_style, company_slug, jd)
    for label in ("primary","cv"):
        block = result[label]
        print(block["docx"]); print(block["report"]); print(block["lint"])

if __name__ == "__main__":
    main()
